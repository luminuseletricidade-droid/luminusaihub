"""
Multi-Format File Extractor
Supports extraction from: PDF, CSV, PNG, JPG, Excel, Doc
Uses Gemini Vision for images and specialized libraries for other formats.
"""

import os
import csv
import json
import logging
import tempfile
import mimetypes
from typing import Dict, Any, Optional, Tuple
from pathlib import Path
import base64

logger = logging.getLogger(__name__)


class MultiFormatExtractor:
    """
    Extracts text and data from multiple file formats.
    Routes to appropriate extraction method based on file type.
    """

    # Supported extensions mapped to extraction methods
    SUPPORTED_FORMATS = {
        # PDF
        '.pdf': 'extract_pdf',
        # Images (will use Gemini Vision)
        '.png': 'extract_image',
        '.jpg': 'extract_image',
        '.jpeg': 'extract_image',
        '.gif': 'extract_image',
        '.webp': 'extract_image',
        '.bmp': 'extract_image',
        # Excel
        '.xlsx': 'extract_excel',
        '.xls': 'extract_excel',
        # Word Documents
        '.docx': 'extract_docx',
        '.doc': 'extract_doc_legacy',
        # CSV/Text
        '.csv': 'extract_csv',
        '.txt': 'extract_text',
        # JSON
        '.json': 'extract_json',
    }

    def __init__(self):
        """Initialize the extractor with optional Gemini client for images."""
        self.gemini_client = None
        self._init_gemini()

    def _init_gemini(self):
        """Initialize Gemini client for image extraction."""
        try:
            gemini_api_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
            if gemini_api_key:
                from google import genai
                self.gemini_client = genai.Client(api_key=gemini_api_key)
                logger.info("Gemini client initialized for image extraction")
            else:
                logger.warning("No Gemini API key found - image extraction will be limited")
        except Exception as e:
            logger.warning(f"Could not initialize Gemini client: {e}")

    def get_file_extension(self, filename: str) -> str:
        """Get lowercase file extension from filename."""
        return Path(filename).suffix.lower()

    def is_supported(self, filename: str) -> bool:
        """Check if file format is supported."""
        ext = self.get_file_extension(filename)
        return ext in self.SUPPORTED_FORMATS

    def get_mime_type(self, filename: str) -> str:
        """Get MIME type for a file."""
        mime_type, _ = mimetypes.guess_type(filename)
        return mime_type or 'application/octet-stream'

    def extract(
        self,
        file_content: bytes,
        filename: str,
        category: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Extract text/data from a file.

        Args:
            file_content: Raw file bytes
            filename: Original filename (used to detect format)
            category: Optional category for context

        Returns:
            Dictionary with:
                - success: bool
                - text: extracted text content
                - data: structured data if applicable (CSV, Excel, JSON)
                - method: extraction method used
                - error: error message if failed
        """
        ext = self.get_file_extension(filename)

        if ext not in self.SUPPORTED_FORMATS:
            return {
                "success": False,
                "text": None,
                "data": None,
                "method": None,
                "error": f"Unsupported file format: {ext}"
            }

        method_name = self.SUPPORTED_FORMATS[ext]
        method = getattr(self, method_name, None)

        if not method:
            return {
                "success": False,
                "text": None,
                "data": None,
                "method": method_name,
                "error": f"Extraction method not implemented: {method_name}"
            }

        try:
            logger.info(f"Extracting {filename} using method: {method_name}")
            result = method(file_content, filename)
            result["method"] = method_name
            return result
        except Exception as e:
            logger.error(f"Error extracting {filename}: {e}")
            return {
                "success": False,
                "text": None,
                "data": None,
                "method": method_name,
                "error": str(e)
            }

    def extract_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from PDF using Gemini Vision OCR.

        Converts PDF pages to images and uses Gemini Vision to extract text.
        This approach is more reliable than traditional text extraction,
        especially for scanned documents.
        """
        # If Gemini is not available, fall back to basic extraction
        if not self.gemini_client:
            logger.warning("Gemini client not available, using basic PDF extraction")
            return self._extract_pdf_basic(file_content)

        try:
            from pdf2image import convert_from_bytes
            from google.genai import types
            import io

            # Convert PDF to images - extract ALL pages for complete analysis
            # Gemini LLM can handle large text and extract relevant identifiers
            try:
                images = convert_from_bytes(
                    file_content,
                    dpi=200,  # Good quality for OCR accuracy
                    first_page=1,
                    last_page=10  # Up to 10 pages for thorough extraction
                )
            except Exception as conv_error:
                logger.warning(f"pdf2image conversion failed: {conv_error}, trying basic extraction")
                return self._extract_pdf_basic(file_content)

            if not images:
                logger.warning("No images extracted from PDF, trying basic extraction")
                return self._extract_pdf_basic(file_content)

            all_text = []
            extraction_prompt = """Extraia TODO o texto desta página de documento.
Mantenha a formatação e estrutura original do texto.
Se houver tabelas, extraia os dados preservando a estrutura.
Retorne APENAS o texto extraído, sem comentários adicionais."""

            for i, image in enumerate(images):
                try:
                    # Convert PIL Image to JPEG bytes
                    img_buffer = io.BytesIO()
                    image.save(img_buffer, format='JPEG', quality=85)
                    img_bytes = img_buffer.getvalue()

                    # Send to Gemini Vision for OCR
                    response = self.gemini_client.models.generate_content(
                        model="gemini-2.5-flash",
                        contents=[
                            types.Content(
                                role="user",
                                parts=[
                                    types.Part(text=extraction_prompt),
                                    types.Part(
                                        inline_data=types.Blob(
                                            mime_type="image/jpeg",
                                            data=img_bytes
                                        )
                                    )
                                ]
                            )
                        ],
                        config=types.GenerateContentConfig(
                            temperature=0.1,
                            max_output_tokens=8192
                        )
                    )

                    page_text = response.text.strip()
                    if page_text and page_text != "[Imagem sem texto legível]":
                        all_text.append(f"--- Página {i+1} ---\n{page_text}")

                except Exception as page_error:
                    logger.warning(f"Error extracting page {i+1}: {page_error}")
                    continue

            if not all_text:
                # If Gemini Vision failed for all pages, try basic extraction
                logger.warning("Gemini Vision extracted no text, trying basic extraction")
                return self._extract_pdf_basic(file_content)

            combined_text = "\n\n".join(all_text)
            logger.info(f"Gemini Vision OCR extracted {len(combined_text)} chars from {len(images)} pages")

            return {
                "success": True,
                "text": combined_text,
                "data": {
                    "pages": len(images),
                    "method": "gemini_vision_ocr",
                    "has_tables": False  # Could be detected later
                },
                "error": None
            }

        except ImportError as e:
            logger.warning(f"pdf2image not available: {e}, using basic extraction")
            return self._extract_pdf_basic(file_content)
        except Exception as e:
            logger.error(f"Gemini Vision OCR failed: {e}")
            # Fallback to basic PDF extraction
            return self._extract_pdf_basic(file_content)

    def _extract_pdf_basic(self, file_content: bytes) -> Dict[str, Any]:
        """Basic PDF extraction fallback."""
        try:
            import pdfplumber
            import io

            text_parts = []
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            return {
                "success": True,
                "text": "\n\n".join(text_parts),
                "data": {"pages": len(text_parts)},
                "error": None
            }
        except Exception as e:
            return {
                "success": False,
                "text": None,
                "data": None,
                "error": f"PDF extraction failed: {e}"
            }

    def extract_image(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from images using Gemini Vision.
        Can read text, tables, and structured data from images.
        """
        if not self.gemini_client:
            return {
                "success": False,
                "text": None,
                "data": None,
                "error": "Gemini client not available for image extraction"
            }

        try:
            from google.genai import types

            # Encode image to base64
            image_base64 = base64.b64encode(file_content).decode('utf-8')
            mime_type = self.get_mime_type(filename)

            # Create prompt for text extraction
            extraction_prompt = """Analise esta imagem e extraia TODO o texto visível.

Se a imagem contiver:
- Documento: extraia o texto completo mantendo a estrutura
- Tabela: extraia os dados em formato estruturado
- Formulário: extraia campos e valores
- Texto manuscrito: transcreva o melhor possível

Retorne APENAS o texto extraído, sem comentários adicionais.
Se não houver texto legível, responda: "[Imagem sem texto legível]"
"""

            # Call Gemini Vision
            response = self.gemini_client.models.generate_content(
                model="gemini-2.5-flash",
                contents=[
                    types.Content(
                        role="user",
                        parts=[
                            types.Part(text=extraction_prompt),
                            types.Part(
                                inline_data=types.Blob(
                                    mime_type=mime_type,
                                    data=file_content
                                )
                            )
                        ]
                    )
                ],
                config=types.GenerateContentConfig(
                    temperature=0.1,
                    max_output_tokens=8192
                )
            )

            extracted_text = response.text.strip()

            return {
                "success": True,
                "text": extracted_text,
                "data": {
                    "source": "gemini_vision",
                    "mime_type": mime_type
                },
                "error": None
            }

        except Exception as e:
            logger.error(f"Image extraction error: {e}")
            return {
                "success": False,
                "text": None,
                "data": None,
                "error": f"Image extraction failed: {e}"
            }

    def extract_excel(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract data from Excel files (.xlsx, .xls)."""
        try:
            import io
            ext = self.get_file_extension(filename)

            # Try openpyxl for .xlsx
            if ext == '.xlsx':
                return self._extract_xlsx(file_content)
            else:
                # Try xlrd for .xls or fallback
                return self._extract_xls(file_content)

        except ImportError as e:
            return {
                "success": False,
                "text": None,
                "data": None,
                "error": f"Excel library not installed: {e}"
            }
        except Exception as e:
            return {
                "success": False,
                "text": None,
                "data": None,
                "error": f"Excel extraction failed: {e}"
            }

    def _extract_xlsx(self, file_content: bytes) -> Dict[str, Any]:
        """Extract from .xlsx using openpyxl."""
        try:
            from openpyxl import load_workbook
            import io

            wb = load_workbook(io.BytesIO(file_content), data_only=True)

            all_text = []
            all_data = {}

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_data = []
                sheet_text = [f"=== Planilha: {sheet_name} ==="]

                for row in sheet.iter_rows(values_only=True):
                    # Skip completely empty rows
                    if not any(cell is not None for cell in row):
                        continue

                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    sheet_data.append(row_values)
                    sheet_text.append(" | ".join(row_values))

                all_data[sheet_name] = sheet_data
                all_text.extend(sheet_text)
                all_text.append("")  # Blank line between sheets

            return {
                "success": True,
                "text": "\n".join(all_text),
                "data": {
                    "sheets": all_data,
                    "sheet_count": len(wb.sheetnames)
                },
                "error": None
            }
        except ImportError:
            return {
                "success": False,
                "text": None,
                "data": None,
                "error": "openpyxl not installed. Run: pip install openpyxl"
            }

    def _extract_xls(self, file_content: bytes) -> Dict[str, Any]:
        """Extract from .xls using xlrd."""
        try:
            import xlrd
            import io

            wb = xlrd.open_workbook(file_contents=file_content)

            all_text = []
            all_data = {}

            for sheet_name in wb.sheet_names():
                sheet = wb.sheet_by_name(sheet_name)
                sheet_data = []
                sheet_text = [f"=== Planilha: {sheet_name} ==="]

                for row_idx in range(sheet.nrows):
                    row_values = [str(sheet.cell_value(row_idx, col_idx))
                                  for col_idx in range(sheet.ncols)]
                    sheet_data.append(row_values)
                    sheet_text.append(" | ".join(row_values))

                all_data[sheet_name] = sheet_data
                all_text.extend(sheet_text)
                all_text.append("")

            return {
                "success": True,
                "text": "\n".join(all_text),
                "data": {
                    "sheets": all_data,
                    "sheet_count": len(wb.sheet_names())
                },
                "error": None
            }
        except ImportError:
            return {
                "success": False,
                "text": None,
                "data": None,
                "error": "xlrd not installed. Run: pip install xlrd"
            }

    def extract_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from .docx files."""
        try:
            from docx import Document
            import io

            doc = Document(io.BytesIO(file_content))

            paragraphs = []
            tables_text = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_rows = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_rows.append(" | ".join(row_text))

                if table_rows:
                    tables_text.append(f"\n=== Tabela {table_idx + 1} ===")
                    tables_text.extend(table_rows)

            full_text = "\n".join(paragraphs)
            if tables_text:
                full_text += "\n" + "\n".join(tables_text)

            return {
                "success": True,
                "text": full_text,
                "data": {
                    "paragraph_count": len(paragraphs),
                    "table_count": len(doc.tables)
                },
                "error": None
            }
        except ImportError:
            return {
                "success": False,
                "text": None,
                "data": None,
                "error": "python-docx not installed. Run: pip install python-docx"
            }
        except Exception as e:
            return {
                "success": False,
                "text": None,
                "data": None,
                "error": f"DOCX extraction failed: {e}"
            }

    def extract_doc_legacy(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from legacy .doc files.
        Uses Gemini Vision as fallback since .doc parsing is complex.
        """
        # Try to use antiword or similar if available
        try:
            import subprocess
            import tempfile

            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name

            try:
                # Try antiword (common on Linux)
                result = subprocess.run(
                    ['antiword', tmp_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0:
                    return {
                        "success": True,
                        "text": result.stdout,
                        "data": {"method": "antiword"},
                        "error": None
                    }
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass
            finally:
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)

        except Exception as e:
            logger.warning(f"antiword extraction failed: {e}")

        # Fallback: Return warning that .doc is not fully supported
        return {
            "success": False,
            "text": None,
            "data": None,
            "error": "Legacy .doc format not fully supported. Please convert to .docx"
        }

    def extract_csv(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract data from CSV files."""
        try:
            import io

            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    text_content = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                return {
                    "success": False,
                    "text": None,
                    "data": None,
                    "error": "Could not decode CSV file"
                }

            # Parse CSV
            reader = csv.reader(io.StringIO(text_content))
            rows = list(reader)

            # Generate text representation
            text_lines = []
            for row in rows:
                text_lines.append(" | ".join(row))

            return {
                "success": True,
                "text": "\n".join(text_lines),
                "data": {
                    "rows": rows,
                    "row_count": len(rows),
                    "column_count": len(rows[0]) if rows else 0
                },
                "error": None
            }
        except Exception as e:
            return {
                "success": False,
                "text": None,
                "data": None,
                "error": f"CSV extraction failed: {e}"
            }

    def extract_text(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from plain text files."""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    text = file_content.decode(encoding)
                    return {
                        "success": True,
                        "text": text,
                        "data": {"encoding": encoding},
                        "error": None
                    }
                except UnicodeDecodeError:
                    continue

            return {
                "success": False,
                "text": None,
                "data": None,
                "error": "Could not decode text file"
            }
        except Exception as e:
            return {
                "success": False,
                "text": None,
                "data": None,
                "error": f"Text extraction failed: {e}"
            }

    def extract_json(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract data from JSON files."""
        try:
            text = file_content.decode('utf-8')
            data = json.loads(text)

            # Create readable text representation
            pretty_text = json.dumps(data, indent=2, ensure_ascii=False)

            return {
                "success": True,
                "text": pretty_text,
                "data": data,
                "error": None
            }
        except json.JSONDecodeError as e:
            return {
                "success": False,
                "text": None,
                "data": None,
                "error": f"Invalid JSON: {e}"
            }
        except Exception as e:
            return {
                "success": False,
                "text": None,
                "data": None,
                "error": f"JSON extraction failed: {e}"
            }


# Convenience function for quick extraction
def extract_file(
    file_content: bytes,
    filename: str,
    category: Optional[str] = None
) -> Dict[str, Any]:
    """
    Convenience function to extract content from any supported file.

    Args:
        file_content: Raw file bytes
        filename: Original filename
        category: Optional category for context

    Returns:
        Extraction result dictionary
    """
    extractor = MultiFormatExtractor()
    return extractor.extract(file_content, filename, category)


def is_supported_format(filename: str) -> bool:
    """Check if a filename has a supported extension."""
    ext = Path(filename).suffix.lower()
    return ext in MultiFormatExtractor.SUPPORTED_FORMATS


def get_supported_formats() -> list:
    """Return list of supported file extensions."""
    return list(MultiFormatExtractor.SUPPORTED_FORMATS.keys())
